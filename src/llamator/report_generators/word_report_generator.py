import json
import logging
import os

import pandas as pd
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ==========================
# Global Constants
# ==========================
JSON_FILE_PATH = os.path.join(os.path.dirname(__file__), "..", "attacks", "attack_descriptions.json")


# ==========================
# Helper Functions
# ==========================


def set_table_border(table, border_color="404040", border_size=4, border_space=0, border_type="single"):
    """
    Sets gray-black borders for the table.

    Parameters
    ----------
    table : docx.table.Table
        The table to which borders will be applied.
    border_color : str, optional
        The border color in hexadecimal format (default is "404040" - dark gray).
    border_size : int, optional
        The border size (default is 4).
    border_space : int, optional
        The space between the text and the border (default is 0).
    border_type : str, optional
        The border type (e.g., 'single', 'dashed'; default is 'single').
    """
    tbl = table._element
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")

    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        edge_element = OxmlElement(f"w:{edge}")
        edge_element.set(qn("w:val"), border_type)
        edge_element.set(qn("w:sz"), str(border_size))
        edge_element.set(qn("w:space"), str(border_space))
        edge_element.set(qn("w:color"), border_color)
        tblBorders.append(edge_element)

    tblPr.append(tblBorders)


def set_page_background(document, color):
    """
    Sets a pale blue background for the entire document, mimicking stationery paper.

    Parameters
    ----------
    document : docx.Document
        The document for which the background will be set.
    color : str
        The background color in hexadecimal format (e.g., "F0F8FF" for AliceBlue).
    """
    background = OxmlElement("w:background")
    background.set(qn("w:color"), color)
    document.element.insert(0, background)


def load_tests_json(json_file_path: str) -> dict:
    """
    Loads a JSON file containing test information and returns a mapping dictionary.

    Parameters
    ----------
    json_file_path : str
        The path to the JSON file.

    Returns
    -------
    dict
        A dictionary mapping 'in_code_name' to the corresponding test data.
    """
    try:
        with open(json_file_path, encoding="utf-8") as f:
            tests = json.load(f)
        # Create a mapping: in_code_name -> test data
        tests_mapping = {test["in_code_name"]: test for test in tests}
        return tests_mapping
    except Exception as e:
        logging.error(f"Error loading JSON file {json_file_path}: {e}")
        return {}


def set_cell_background(cell, fill_color):
    """
    Sets the background color of a table cell.

    Parameters
    ----------
    cell : docx.table._Cell
        The table cell to color.
    fill_color : str
        The fill color in hexadecimal format (e.g., "FFDAB9").
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_color)
    tcPr.append(shd)


def set_table_background(table, fill_color):
    """
    Sets the background color of the entire table.

    Parameters
    ----------
    table : docx.table.Table
        The table to color.
    fill_color : str
        The fill color in hexadecimal format (e.g., "DCEBFF").
    """
    for row in table.rows:
        for cell in row.cells:
            set_cell_background(cell, fill_color)


# ==========================
# Main Report Generation Function
# ==========================


def create_word_report(
    artifacts_dir: str,
    csv_folder_name: str = "csv_report",
    docx_file_name: str = "attacks_report.docx",
    status_legend: dict = None,
    main_statuses: tuple = ("Broken", "Resilient", "Errors"),
    language: str = "en",
) -> None:
    """
    Generates a well-formatted Word report based on testing results in the chosen language.

    Parameters
    ----------
    artifacts_dir : str
        The path to the directory where artifacts are stored.
    csv_folder_name : str, optional
        The name of the folder containing CSV files within the artifacts directory (default is "csv_report").
    docx_file_name : str, optional
        The name of the generated DOCX file (default is "attacks_report.docx").
    status_legend : dict, optional
        A dictionary mapping statuses to their descriptions. If None, defaults are used.
    main_statuses : tuple, optional
        A tuple of main statuses to display in the report (default is ("Broken", "Resilient", "Errors")).
    language : str, optional
        The language of the report ("en" for English, "ru" for Russian; default is "en").
    """

    # Define localized strings
    strings = {
        "ru": {
            "framework_title": "LLAMATOR",
            "testing_report": "Отчёт о тестировании",
            "legend_title": "Легенда статусов:",
            "status": "Статус",
            "attempts": "Количество попыток",
            "attack_results": "Результаты атаки: {}",
            "report_created": "Отчёт в Word создан: {}",
            "report_failed": "Не удалось создать отчёт в Word: {}",
            "csv_not_found": "Папка с CSV-файлами не найдена: {}",
            "csv_read_failed": "Не удалось прочитать CSV-файл {}: {}",
            "status_column_missing": "Столбец 'status' не найден в {}. Пропуск файла.",
            "test_not_found": "Тест с `in_code_name = '{in_code}'` не найден в JSON. Пропуск файла {file}.",
            "description_missing": "Описание для теста '{test}' не найдено. Используется название CSV файла.",
            "total": "Итого",
        },
        "en": {
            "framework_title": "LLAMATOR",
            "testing_report": "Testing Report",
            "legend_title": "Status Legend:",
            "status": "Status",
            "attempts": "Number of Attempts",
            "attack_results": "Attack Results: {}",
            "report_created": "Word report created: {}",
            "report_failed": "Failed to create Word report: {}",
            "csv_not_found": "CSV folder not found: {}",
            "csv_read_failed": "Failed to read CSV file {}: {}",
            "status_column_missing": "The 'status' column is missing in {}. Skipping file.",
            "test_not_found": "Test with `in_code_name = '{in_code}'` not found in JSON. Skipping file {file}.",
            "description_missing": "Description for test '{test}' not found. Using CSV file name.",
            "total": "Total",
        },
    }

    lang = strings.get(language, strings["en"])  # Default to English if language key is not found

    # Load JSON data with test information
    tests_mapping = load_tests_json(JSON_FILE_PATH)

    if not tests_mapping:
        logging.error("Failed to load test information. Aborting report generation.")
        return

    # Define default status legend if not provided
    if status_legend is None:
        if language == "ru":
            status_legend = {
                "Broken": "Модель не смогла противостоять атаке.",
                "Resilient": "Модель успешно отразила атаку.",
                "Errors": "Во время выполнения атаки произошла ошибка.",
            }
        elif language == "en":
            status_legend = {
                "Broken": "The model failed to withstand the attack.",
                "Resilient": "The model successfully resisted the attack.",
                "Errors": "An error occurred during the attack attempt.",
            }

    # Full path to the CSV folder
    csv_folder_path = os.path.join(artifacts_dir, csv_folder_name)

    # Full path to the output DOCX file
    output_docx_path = os.path.join(artifacts_dir, docx_file_name)

    # Initialize Word document
    document = Document()

    # Set page margins
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.7)  # Top margin
        section.bottom_margin = Inches(0.7)  # Bottom margin
        section.left_margin = Inches(0.8)  # Left margin
        section.right_margin = Inches(0.8)  # Right margin

    # Configure styles
    styles = document.styles

    # Add style for CenterTitle (Sans-serif for headings)
    if "CenterTitle" not in styles:
        center_title_style = styles.add_style("CenterTitle", WD_STYLE_TYPE.PARAGRAPH)
        center_title_style.font.name = "Helvetica"  # Sans-serif
        center_title_style.font.size = Pt(20)
        center_title_style.font.bold = True
        center_title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        center_title_style.paragraph_format.space_after = Pt(10)
        # Ensure EastAsia and ASCII fonts are also set to Helvetica
        center_title_style.font.eastasia = "Helvetica"
        center_title_style.font.ascii = "Helvetica"
    else:
        center_title_style = styles["CenterTitle"]

    # Add style for Heading2Custom (Sans-serif for subheadings)
    if "Heading2Custom" not in styles:
        heading2_style = styles.add_style("Heading2Custom", WD_STYLE_TYPE.PARAGRAPH)
        heading2_style.base_style = styles["Heading 2"]
        heading2_style.font.name = "Helvetica"  # Sans-serif
        heading2_style.font.size = Pt(14)
        heading2_style.paragraph_format.space_after = Pt(10)
        # Ensure EastAsia and ASCII fonts are also set to Helvetica
        heading2_style.font.eastasia = "Helvetica"
        heading2_style.font.ascii = "Helvetica"
    else:
        heading2_style = styles["Heading2Custom"]

    # Add style for NormalStyle (Serif for body text)
    if "NormalStyle" not in styles:
        normal_style = styles.add_style("NormalStyle", WD_STYLE_TYPE.PARAGRAPH)
        normal_style.base_style = styles["Normal"]
        normal_style.font.name = "Times New Roman"  # Serif
        normal_style.font.size = Pt(12)
        normal_style.paragraph_format.space_after = Pt(6)
        # Ensure EastAsia and ASCII fonts are also set to Times New Roman
        normal_style.font.eastasia = "Times New Roman"
        normal_style.font.ascii = "Times New Roman"
    else:
        normal_style = styles["NormalStyle"]

    # Set a pale blue background for the document
    set_page_background(document, "F0F8FF")  # "F0F8FF" - AliceBlue

    # ==========================
    # Add Initial Content
    # ==========================

    # Add the title "LLAMATOR"
    framework_title = document.add_paragraph(lang["framework_title"], style="CenterTitle")
    for run in framework_title.runs:
        run.font.color.rgb = RGBColor(80, 80, 80)  # Dark gray text color

    # Add a horizontal line
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Create a horizontal line using the bottom border of the paragraph
    p_para = p._element
    p_pPr = p_para.get_or_add_pPr()
    p_pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")  # Line thickness
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")  # Black color
    p_pBdr.append(bottom)
    p_pPr.append(p_pBdr)

    # Add the "Testing Report" title (localized)
    testing_report_title = document.add_paragraph(lang["testing_report"], style="CenterTitle")
    for run in testing_report_title.runs:
        run.font.color.rgb = RGBColor(80, 80, 80)  # Dark gray text color

    document.add_paragraph()  # Add an empty paragraph for spacing

    # Add the status legend title
    legend_title = document.add_paragraph(lang["legend_title"], style="Heading2Custom")
    for run in legend_title.runs:
        run.font.color.rgb = RGBColor(80, 80, 80)  # Dark gray text color

    # Add each status and its description to the legend
    for status, description in status_legend.items():
        p = document.add_paragraph(style="NormalStyle")
        run_status = p.add_run(f"{status}: ")
        run_status.bold = True
        run_status.font.color.rgb = RGBColor(80, 80, 80)  # Dark gray text color
        run_description = p.add_run(description)
        run_description.font.color.rgb = RGBColor(80, 80, 80)  # Dark gray text color

    document.add_paragraph()  # Add an empty paragraph for spacing

    # Check if the CSV folder exists
    if not os.path.isdir(csv_folder_path):
        logging.error(lang["csv_not_found"].format(csv_folder_path))
        return

    # Collect all CSV files into a list
    csv_files = sorted([f for f in os.listdir(csv_folder_path) if f.endswith(".csv")])

    # Iterate through all CSV files to create tables with results
    for idx, csv_file in enumerate(csv_files):
        # Extract in_code_name from the file name
        in_code_name = os.path.splitext(csv_file)[0]
        csv_path = os.path.join(csv_folder_path, csv_file)

        # Get test information from JSON
        test_info = tests_mapping.get(in_code_name)

        if not test_info:
            logging.warning(lang["test_not_found"].format(in_code=in_code_name, file=csv_file))
            # Use the CSV file name as the attack name
            test_name = in_code_name
            test_description = ""
        else:
            test_name = test_info["name"]
            test_description = test_info.get(f"description_{language}")
            if not test_description:
                logging.warning(lang["description_missing"].format(test=test_name))
                # Use the CSV file name as the attack name if description is missing
                test_name = in_code_name
                test_description = ""

        try:
            df = pd.read_csv(csv_path)
        except Exception as e:
            logging.error(lang["csv_read_failed"].format(csv_path, e))
            continue

        # Check for the 'status' column
        if "status" not in df.columns:
            logging.warning(lang["status_column_missing"].format(csv_path))
            continue

        # Aggregate the number of attempts by status
        summary = df["status"].value_counts().to_dict()

        # Ensure all main statuses are present, even if their count is zero
        summary_complete = {status: summary.get(status, 0) for status in main_statuses}

        # Calculate total attempts
        total_attempts = sum(summary_complete.values())

        # Add the attack name
        attack_title = document.add_paragraph(test_name, style="Heading2Custom")
        for run in attack_title.runs:
            run.font.color.rgb = RGBColor(80, 80, 80)  # Dark gray text color

        # Keep the attack title with the next paragraph/table
        attack_title.paragraph_format.keep_with_next = True
        attack_title.paragraph_format.keep_together = True

        # Add the attack description if available
        if test_description:
            attack_description = document.add_paragraph(test_description, style="NormalStyle")
            for run in attack_description.runs:
                run.font.color.rgb = RGBColor(80, 80, 80)  # Dark gray text color

            # Keep the attack description with the next table
            attack_description.paragraph_format.keep_with_next = True
            attack_description.paragraph_format.keep_together = True

        # Add a table to the document
        table = document.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        table.columns[0].width = Inches(3)
        table.columns[1].width = Inches(2)

        # Add table headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = lang["status"]
        hdr_cells[1].text = lang["attempts"]

        # Format table headers
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(80, 80, 80)  # Dark gray text color
                    run.font.name = "Helvetica"  # Sans-serif
                    run.font.eastasia = "Helvetica"
                    run.font.ascii = "Helvetica"
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add data rows to the table
        for status, count in summary_complete.items():
            row_cells = table.add_row().cells
            row_cells[0].text = status
            row_cells[1].text = str(count)
            # Align text to center and set color
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(12)
                        run.font.color.rgb = RGBColor(80, 80, 80)  # Dark gray text color
                        run.font.name = "Times New Roman"  # Serif
                        run.font.eastasia = "Times New Roman"
                        run.font.ascii = "Times New Roman"

            # Apply specific cell shading based on status
            if status == "Resilient" and language == "en" or status == "Resilient" and language == "ru":
                # Pale Green
                set_cell_background(row_cells[0], "DFFFD6")
            elif status == "Broken" and language == "en" or status == "Broken" and language == "ru":
                # Pale Orange
                set_cell_background(row_cells[0], "FFDAB9")
            elif status == "Errors" and language == "en" or status == "Errors" and language == "ru":
                # Pale Yellow
                set_cell_background(row_cells[0], "FFFFE0")

        # Add the total row
        total_row = table.add_row().cells
        total_label = lang["total"]
        total_row[0].text = total_label
        total_row[1].text = str(total_attempts)
        # Format the total row
        for cell in total_row:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(80, 80, 80)  # Dark gray text color
                    run.font.name = "Times New Roman"  # Serif
                    run.font.eastasia = "Times New Roman"
                    run.font.ascii = "Times New Roman"

        # Set a light gray background for the cells of the summary row
        set_cell_background(total_row[0], "d6d6d6")  # Light gray background
        set_cell_background(total_row[1], "d6d6d6")  # Light gray background

        # Set table borders
        set_table_border(table, border_color="404040", border_size=4, border_space=0, border_type="single")

        # Change the fill of the entire table to a lighter shade
        set_table_background(table, "ededed")

        # Prevent table rows from breaking across pages
        for row in table.rows:
            row.allow_break_across_pages = False

        # Add a blank paragraph for indentation only if it is not the last CSV file
        if idx < len(csv_files) - 1:
            document.add_paragraph()  # Add a blank paragraph between tables

    # Save the document
    try:
        document.save(output_docx_path)
        logging.info(lang["report_created"].format(output_docx_path))
        print(f"Word report created: {output_docx_path}")
    except Exception as e:
        logging.error(lang["report_failed"].format(e))
